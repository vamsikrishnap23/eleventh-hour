from flask import Flask, request, render_template, flash, redirect, url_for, session, send_from_directory, send_file
import os
import shutil
import json
import subprocess
import traceback
import io
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph
from reportlab.lib.units import inch
from PIL import Image
from dotenv import load_dotenv
import google.generativeai as genai
import re
import docx

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = './uploads'
app.config['ALLOWED_EXTENSIONS'] = {'.py', '.js', '.jsx', '.ts', '.tsx', '.css', '.java', '.c', '.cpp', '.h', '.cs', '.go', '.rb', '.php', '.html'}
app.secret_key = 'super_secret_key'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure Gemini API
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

def allowed_file(filename):
    return os.path.splitext(filename)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def analyze_code(code, filename):
    try:
        model = genai.GenerativeModel('gemini-1.5-pro')
        prompt = f"""You are a code analysis assistant. Analyze the following code from the file `{filename}`:

```{os.path.splitext(filename)[1][1:]}
{code}
```

Respond with a raw JSON object only, no markdown or commentary.

Mermaid Format Notes:
- Use simple, alphanumeric node labels only (e.g., A, B, C or Step1, Step2).
- Avoid special characters like "=", ".", "()", etc. inside node labels.
- Use square brackets for all nodes.
- Use TD direction (top-down).
- One connection per line.

JSON Format:
{{ "summary": "...", "mermaid": "graph TD\\nA[Start]\\nB[Process]\\nC[End]\\nA-->B\\nB-->C", "description": "..." }}"""

        response = model.generate_content(prompt)
        raw_response = response.text

        # Debug raw LLM response
        print("\n=== RAW LLM RESPONSE START ===")
        print(raw_response)
        print("=== RAW LLM RESPONSE END ===\n")

        clean_response = raw_response.strip()
        if clean_response.startswith('```json'):
            clean_response = clean_response[7:]
        if clean_response.endswith('```'):
            clean_response = clean_response[:-3]
        clean_response = clean_response.strip()

        data = json.loads(clean_response)
        print(f"Parsed JSON for {filename}: {data}")

        if not data.get('mermaid') or 'graph' not in data['mermaid'] or '-->' not in data['mermaid']:
            data['mermaid'] = None
            data['description'] = f"Fallback: Workflow for {filename}"
        else:
            data['description'] = data.get('description', f"Flowchart for {filename}")
            try:
                # Convert Mermaid code to PNG using mermaid-cli
                mermaid_code = data['mermaid'].replace('\\n', '\n')
                
                # Clean up node labels to ensure proper syntax
                def clean_mermaid_code(code):
                    # Split into lines and clean each line
                    lines = code.split('\n')
                    cleaned_lines = []
                    
                    for line in lines:
                        # Skip empty lines
                        if not line.strip():
                            continue
                        
                        # Keep graph TD line as is
                        if line.startswith('graph TD'):
                            cleaned_lines.append(line)
                            continue
                            
                        # Clean node definitions and connections
                        line = line.strip()
                        # Remove semicolons
                        line = line.replace(';', '')
                        # Replace curly braces with square brackets
                        line = re.sub(r'\{([^}]*)\}', r'[\1]', line)
                        
                        # Clean up node labels
                        def clean_label(match):
                            label = match.group(1)
                            # Remove characters that confuse Mermaid
                            label = re.sub(r'[=\[\]{}"\'`<>;]+', '', label)
                            label = label.strip()
                            label = label.replace(' ', '_')  # Optional: Replace spaces with underscores
                            return f'[{label}]'
                        
                        line = re.sub(r'\[(.*?)\]', clean_label, line)
                        cleaned_lines.append(line)
                    
                    return '\n'.join(cleaned_lines)
                
                mermaid_code = clean_mermaid_code(mermaid_code)
                print(f"\nCleaned Mermaid code for {filename}:\n{mermaid_code}\n")
                
                mmd_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{os.path.splitext(filename)[0]}.mmd")
                with open(mmd_path, 'w') as f:
                    f.write(mermaid_code)

                png_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{os.path.splitext(filename)[0]}_flowchart.png")
                if shutil.which("mmdc") is None:
                    raise EnvironmentError("Mermaid CLI (mmdc) is not installed or not in PATH")

                result = subprocess.run([
                    "mmdc", "-i", mmd_path, "-o", png_path, "-t", "default"
                ], capture_output=True, text=True)

                if result.returncode != 0:
                    raise RuntimeError(f"Mermaid CLI error: {result.stderr}")

                data['flowchart_path'] = png_path
            except Exception as e:
                print(f"Flowchart generation failed for {filename}: {e}")
                print(traceback.format_exc())
                data['mermaid'] = None
                data['description'] = f"Failed to generate flowchart for {filename}"

        return data

    except Exception as e:
        print(f"Error analyzing {filename}: {e}")
        return {
            'summary': f"Could not analyze {filename}",
            'mermaid': None,
            'description': f"Fallback: Workflow for {filename}"
        }

def generate_pdf(analysis):
    try:
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        style = ParagraphStyle(
            'Custom',
            parent=styles['Normal'],
            fontSize=12,
            leading=16,
            spaceAfter=20
        )
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Title'],
            fontSize=24,
            leading=30,
            spaceAfter=30,
            alignment=1
        )

        # Title page
        title = Paragraph("Code Documentation", title_style)
        title.wrapOn(p, 400, 100)
        title.drawOn(p, 100, 700)
        
        y_position = 650

        # Add author details if available
        author_details = session.get('author_details', {})
        if author_details:
            author_name = author_details.get('name', 'Unknown Author')
            author_text = f"<b>Author:</b> {author_name}"
            author_para = Paragraph(author_text, style)
            author_para.wrapOn(p, 400, 100)
            author_para.drawOn(p, 100, y_position)
            y_position -= 30

            # Remove caption prefix section
            y_position -= 50
        else:
            fallback_text = "Author details not provided."
            fallback_para = Paragraph(fallback_text, style)
            fallback_para.wrapOn(p, 400, 100)
            fallback_para.drawOn(p, 100, y_position)
            y_position -= 30

        # Add project details if available
        project_details = session.get('project_details', {})
        if project_details:
            p.setFont("Helvetica-Bold", 14)
            p.drawString(100, y_position, "Project Details")
            y_position -= 30
            
            goal_text = f"<b>Project Goal:</b><br/>{project_details.get('goal', 'N/A')}"
            goal_para = Paragraph(goal_text, style)
            w, h = goal_para.wrapOn(p, 400, 200)
            goal_para.drawOn(p, 100, y_position - h)
            y_position -= (h + 30)
            
            audience_text = f"<b>Target Audience:</b><br/>{project_details.get('audience', 'N/A')}"
            audience_para = Paragraph(audience_text, style)
            w, h = audience_para.wrapOn(p, 400, 200)
            audience_para.drawOn(p, 100, y_position - h)
        else:
            fallback_text = "Project details not provided."
            fallback_para = Paragraph(fallback_text, style)
            fallback_para.wrapOn(p, 400, 100)
            fallback_para.drawOn(p, 100, y_position - 30)

        p.showPage()

        # Analysis pages
        for filename, data in analysis.items():
            y_position = 750
            header_text = f"Analysis for {filename}"
            header_para = Paragraph(header_text, ParagraphStyle('Header', parent=style, fontSize=16, spaceAfter=20))
            w, h = header_para.wrapOn(p, 400, 100)
            header_para.drawOn(p, 100, y_position - h)
            y_position -= (h + 20)
            
            summary_text = f"<b>Summary:</b><br/>{data['summary']}"
            summary_para = Paragraph(summary_text, style)
            w, h = summary_para.wrapOn(p, 400, 200)
            summary_para.drawOn(p, 100, y_position - h)
            y_position -= (h + 20)
            
            desc_text = f"<b>Description:</b><br/>{data['description']}"
            desc_para = Paragraph(desc_text, style)
            w, h = desc_para.wrapOn(p, 400, 200)
            desc_para.drawOn(p, 100, y_position - h)
            y_position -= (h + 20)
            
            if data.get('flowchart_path') and os.path.exists(data['flowchart_path']):
                try:
                    img = Image.open(data['flowchart_path'])
                    width, height = img.size
                    page_width = 400
                    scale_factor = page_width / width
                    new_width = page_width
                    new_height = int(height * scale_factor)
                    max_height = y_position - 50
                    if new_height > max_height:
                        scale_factor = max_height / height
                        new_width = int(width * scale_factor)
                        new_height = max_height
                    
                    img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{os.path.basename(data['flowchart_path'])}")
                    img.save(temp_path, quality=100, optimize=False)
                    
                    x_position = (letter[0] - new_width) / 2
                    caption_prefix = session.get('author_details', {}).get('caption_prefix', 'Fig.')
                    figure_num = list(analysis.keys()).index(filename) + 1
                    caption_text = f"{caption_prefix} {figure_num}: Flowchart for {filename}"
                    caption_para = Paragraph(caption_text, style)
                    w, h = caption_para.wrapOn(p, 400, 100)
                    caption_para.drawOn(p, x_position, y_position - new_height - h - 10)
                    p.drawImage(temp_path, x_position, y_position - new_height, width=new_width, height=new_height, preserveAspectRatio=True)
                    os.remove(temp_path)
                except Exception as e:
                    print(f"Error adding flowchart to PDF: {e}")
                    error_para = Paragraph("Error: Could not add flowchart to PDF", style)
                    error_para.wrapOn(p, 400, 100)
                    error_para.drawOn(p, 100, y_position - 50)
            
            # Page number
            p.setFont("Helvetica", 10)
            page_num = len(analysis.keys()) + 1  # Account for title page
            p.drawString(letter[0] - 100, 30, f"Page {page_num}")
            p.showPage()

        p.save()
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Error generating PDF: {e}")
        print(traceback.format_exc())
        raise

def generate_abstract(analysis):
    try:
        model = genai.GenerativeModel('gemini-1.5-pro')
        summary_text = " ".join(data['summary'] for data in analysis.values())
        prompt = f"""Generate a concise academic abstract (100-150 words) summarizing the following code analysis:
        {summary_text}
        The abstract should include the purpose, methodology, key findings, and significance."""
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"Error generating abstract: {e}")
        return "Abstract not available. Please summarize the project manually."

def add_table_of_contents(paragraph):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-4" \\h \\z \\u')  # field code
    paragraph._p.append(fldSimple)

def generate_llm_section(model, prompt):
    try:
        response = model.generate_content(prompt)
        text = response.text.strip()
        print(f"\n=== RAW LLM RESPONSE ===\n{text}\n=== END RAW LLM RESPONSE ===")
        
        # Remove markdown code block markers if present
        if text.startswith('```') and text.endswith('```'):
            text = text[3:-3].strip()
        
        sections = {}
        current_section = None
        current_content = []
        
        for line in text.split('\n'):
            line = line.strip()
            # Check for section headers (both markdown and plain text formats)
            if line.startswith('## ') or line.startswith('### '):
                # If we have a current section, save its content
                if current_section:
                    sections[current_section] = '\n'.join(current_content).strip()
                    current_content = []
                # Extract section name and normalize it
                current_section = line.replace('#', '').strip().lower()
            elif current_section and line:
                current_content.append(line)
        
        # Save the last section's content
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        if not sections:
            raise ValueError("No sections extracted from LLM response")
            
        # Ensure all required sections exist
        required_sections = ['abstract', 'acknowledgement', 'introduction', 'objectives', 
                           'methodology', 'conclusion', 'future scope']
        for section in required_sections:
            if section not in sections:
                sections[section] = f'[Section "{section}" not generated]'
                
        return sections
    except Exception as e:
        if "429" in str(e):
            print(f"Quota exceeded: {e}. Waiting {e.retry_delay.seconds} seconds before fallback...")
            import time
            time.sleep(e.retry_delay.seconds)
        print(f"Error generating LLM section: {e}")
        return {
            'introduction': '[Quota exceeded or error, please add manually]',
            'objectives': '[Quota exceeded or error, please add manually]',
            'methodology': '[Quota exceeded or error, please add manually]',
            'conclusion': '[Quota exceeded or error, please add manually]',
            'future scope': '[Quota exceeded or error, please add manually]',
            'acknowledgement': '[Quota exceeded or error, please add manually]',
            'abstract': '[Quota exceeded or error, please add manually]'
        }

def generate_word_document(analysis):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Initialize Gemini model
    model = genai.GenerativeModel('gemini-1.5-pro')

    # Get project details and author info
    project_details = session.get('project_details', {})
    author_details = session.get('author_details', {})
    goal = project_details.get('goal', 'N/A')
    project_title = project_details.get('title', 'Project Documentation')

    # Single prompt for all sections
    all_prompt = (
        f"Generate the following sections for a formal project report based on the file '{list(analysis.keys())[0]}' with summary: '{analysis[list(analysis.keys())[0]]['summary']}'. "
        f"The project goal is '{goal}' aimed at the audience '{project_details.get('audience', 'N/A')}'. Use code analysis data where relevant. "
        "Structure the response with ### SectionName headers and content below each. Adopt a strictly impersonal, declarative tone suitable for an official report—avoid phrases like 'The analysis reveals', 'Based on the analysis', or any first-person references. Present content as established facts or recommendations without implying a narrator. Generate:\n"
        "- ### Abstract: A concise academic abstract (100-150 words) summarizing the purpose, method (code review of {list(analysis.keys())[0]}'s language/features), and key findings.\n"
        "- ### Acknowledgement: A formal acknowledgement (50-100 words) recognizing contributions from the author '{author_details.get('name', 'Unknown Author')}' and a placeholder for others (e.g., guide, teammates) in an official manner.\n"
        "- ### Introduction: A concise introduction (100-150 words) outlining the purpose and context based on the project goal.\n"
        "- ### Objectives: Concise objectives (50-100 words) detailing the aims to be achieved for the audience.\n"
        "- ### Methodology: A concise methodology (100-150 words) describing the code review process for '{list(analysis.keys())[0]}', focusing on its language (e.g., JavaScript/React) and features (e.g., components, data flow).\n"
        "- ### Conclusion: A concise conclusion (100-150 words) summarizing key findings from the analysis summaries: {', '.join([data['summary'] for data in analysis.values()])}.\n"
        "- ### Future Scope: A concise future scope (100-150 words) proposing enhancements based on the analysis summaries."
    )
    all_sections = generate_llm_section(model, all_prompt)

    # 1. Title Page
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = project_title
    header_para.alignment = 1
    
    para = doc.add_paragraph()
    run = para.add_run(project_title)
    run.bold = True
    run.font.size = docx.shared.Pt(24)
    para.alignment = 1
    
    para = doc.add_paragraph(f"Author: {author_details.get('name', 'Your Name')}")
    para.alignment = 1
    
    para = doc.add_paragraph(f"Registration Number: {author_details.get('reg_number', '[Your Reg Number]')}")
    para.alignment = 1
    
    if author_details.get('teammates'):
        para = doc.add_paragraph("Teammates:")
        para.alignment = 1
        for teammate in author_details['teammates'].split('\n'):
            if teammate.strip():
                para = doc.add_paragraph(teammate.strip())
                para.alignment = 1
    
    para = doc.add_paragraph(f"Institution & Department: {author_details.get('institution_department', '[Your Institution & Department]')}")
    para.alignment = 1
    
    para = doc.add_paragraph(f"Guide: {author_details.get('guide_name', '[Guide Name, Optional]')}")
    para.alignment = 1
    
    para = doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
    para.alignment = 1
    
    doc.add_page_break()

    # 2. Certificate Page
    doc.add_heading('Certificate', level=1)
    authors = author_details.get('name', 'Your Name')
    if author_details.get('teammates'):
        teammate_names = [t.split('-')[0].strip() for t in author_details['teammates'].split('\n') if t.strip()]
        if teammate_names:
            authors += " and " + ", ".join(teammate_names)
    doc.add_paragraph(
        f"This is to certify that the project titled '{project_title}' was completed by {authors} "
        f"under the guidance of {author_details.get('guide_name', '[Guide Name, Optional]')}."
    )
    doc.add_paragraph(f"Place: {author_details.get('city', '[Your City]')}    Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("Signature: __________    Guide Signature: __________")
    doc.add_page_break()

    # 3. Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(all_sections.get('abstract', '[Section not generated]'))

    # 4. Main Report Body
    doc.add_heading('Main Report', level=1)

    # Introduction
    doc.add_heading('Introduction', level=2)
    doc.add_paragraph(all_sections.get('introduction', '[Section not generated]'))

    # Objectives
    doc.add_heading('Objectives', level=2)
    doc.add_paragraph(all_sections.get('objectives', '[Section not generated]'))

    # Methodology
    doc.add_heading('Methodology', level=2)
    doc.add_paragraph(all_sections.get('methodology', '[Section not generated]'))

    # Results & Analysis
    doc.add_heading('Results & Analysis', level=2)
    for i, (filename, data) in enumerate(analysis.items(), 1):
        doc.add_paragraph(f"{i}. {filename}")
        doc.add_paragraph(f"Summary: {data['summary']}")
        doc.add_paragraph(f"Description: {data['description']}")
        if data.get('flowchart_path') and os.path.exists(data['flowchart_path']):
            doc.add_picture(data['flowchart_path'], width=docx.shared.Inches(5))
            caption_prefix = session.get('author_details', {}).get('caption_prefix', 'Fig.')
            doc.add_paragraph(f"{caption_prefix} {i}: Flowchart for {filename}")

    # Screenshots
    doc.add_heading('Screenshots', level=2)
    doc.add_paragraph("Please add screenshots of your project here. Include relevant screenshots that demonstrate the functionality and user interface of your application.")
    doc.add_paragraph("For each screenshot:")
    doc.add_paragraph("1. Add the image using the 'Insert Picture' feature")
    doc.add_paragraph("2. Add a caption below the image using the format: '[Caption Prefix] [Number]: [Description]'")
    doc.add_paragraph("3. Ensure screenshots are clear and properly sized")
    doc.add_paragraph("4. Include screenshots of key features and user interactions")

    # Conclusion
    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(all_sections.get('conclusion', '[Section not generated]'))

    # Future Scope
    doc.add_heading('Future Scope', level=2)
    doc.add_paragraph(all_sections.get('future scope', '[Section not generated]'))

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/')
def home():
    return render_template('upload.html', current_step='upload')

@app.route('/upload', methods=['GET', 'POST'])
def upload_files():
    try:
        if request.method == 'POST':
            # Clean uploads
            if os.path.exists(app.config['UPLOAD_FOLDER']):
                try:
                    # First, remove all files in the directory
                    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                        try:
                            if os.path.isfile(file_path):
                                os.unlink(file_path)
                            elif os.path.isdir(file_path):
                                shutil.rmtree(file_path)
                        except Exception as e:
                            print(f"Error removing {file_path}: {e}")
                    
                    # Then remove the directory itself
                    shutil.rmtree(app.config['UPLOAD_FOLDER'])
                except Exception as e:
                    print(f"Error clearing uploads folder: {e}")
                    flash('Error clearing previous uploads', 'error')
            
            # Create fresh upload directory
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

            files = request.files.getlist('files')
            if not files or all(f.filename == '' for f in files):
                flash('No files selected', 'error')
                return redirect(url_for('upload_files'))

            # Check total size of all files
            total_size = sum(len(f.read()) for f in files)
            for f in files:
                f.seek(0)  # Reset file pointer after reading
                
            if total_size > 10 * 1024 * 1024:  # 10MB limit
                flash('Total file size exceeds 10MB limit', 'error')
                return redirect(url_for('upload_files'))

            if len(files) > 20:  # Limit number of files
                flash('Maximum 20 files allowed per upload', 'error')
                return redirect(url_for('upload_files'))

            file_paths = []
            for file in files:
                if file and allowed_file(file.filename):
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                    file.save(filepath)
                    file_paths.append(file.filename)
                else:
                    allowed = ', '.join(sorted(app.config['ALLOWED_EXTENSIONS']))
                    flash(f'Invalid file: {file.filename}. Allowed: {allowed}', 'error')
                    return redirect(url_for('upload_files'))

            session['file_paths'] = file_paths

            # Analyze files
            analysis = {}
            for filename in file_paths:
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    result = analyze_code(content, filename)
                    analysis[filename] = result

            # Store only essential data in session
            session['analysis'] = {
                filename: {
                    'summary': data['summary'],
                    'description': data['description'],
                    'flowchart_path': data.get('flowchart_path')
                }
                for filename, data in analysis.items()
            }

            return redirect(url_for('analysis'))

        return render_template('upload.html', 
                            current_step='upload',
                            uploaded_files=session.get('file_paths', []))

    except Exception as e:
        print(f"Upload error: {e}")
        flash('Unexpected error during upload.', 'error')
        return redirect(url_for('upload_files'))

@app.route('/export_pdf')
def export_pdf():
    try:
        analysis = session.get('analysis', {})
        if not analysis:
            flash('No analysis data to export. Please upload and analyze files.', 'error')
            return redirect(url_for('upload_files'))
        
        buffer = generate_pdf(analysis)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='code_analysis.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        print(f"Error in export_pdf: {e}")
        flash('Error generating PDF. Please try again.', 'error')
        return redirect(url_for('analysis'))

@app.route('/analysis')
def analysis():
    analysis = session.get('analysis', {})
    if not analysis:
        flash('No analysis data found. Please upload files.', 'error')
        return redirect(url_for('upload_files'))
    
    # Prepare flowcharts for rendering
    for filename, data in analysis.items():
        if data.get('flowchart_path') and os.path.exists(data['flowchart_path']):
            data['flowchart_url'] = f'/uploads/{os.path.basename(data["flowchart_path"])}'
        else:
            data['flowchart_url'] = None
    
    return render_template('analysis.html', current_step='analysis', analysis=analysis)

@app.route('/questions', methods=['GET', 'POST'])
def questions():
    if request.method == 'POST':
        # Validate and collect form data
        project_title = request.form.get('project_title', '').strip()
        author_name = request.form.get('author_name', '').strip()
        author_reg_number = request.form.get('author_reg_number', '').strip()
        teammates = request.form.get('teammates', '').strip()
        institution_department = request.form.get('institution_department', '').strip()
        guide_name = request.form.get('guide_name', '').strip()
        city = request.form.get('city', '').strip()
        caption_prefix = request.form.get('caption_prefix', 'Fig.').strip()
        project_goal = request.form.get('project_goal', '').strip()
        project_audience = request.form.get('project_audience', '').strip()

        # Validation
        required_fields = {
            'Project Title': project_title,
            'Author Name': author_name,
            'Author Registration Number': author_reg_number,
            'Institution & Department': institution_department,
            'City': city,
            'Project Goal': project_goal,
            'Project Audience': project_audience
        }
        for field_name, field_value in required_fields.items():
            if not field_value:
                flash(f'{field_name} is required.', 'error')
                return redirect(url_for('questions'))

        if not caption_prefix.replace('.', '').isalnum():
            flash('Caption Prefix must be alphanumeric or contain only dots.', 'error')
            return redirect(url_for('questions'))

        # Store in session
        session['author_details'] = {
            'name': author_name,
            'reg_number': author_reg_number,
            'teammates': teammates,
            'institution_department': institution_department,
            'guide_name': guide_name,
            'city': city,
            'caption_prefix': caption_prefix
        }
        session['project_details'] = {
            'title': project_title,
            'goal': project_goal,
            'audience': project_audience
        }
        flash('Questions and details saved successfully.', 'success')
        return redirect(url_for('analysis'))

    return render_template('questions.html', current_step='questions')

@app.route('/generate')
def generate_document():
    analysis = session.get('analysis', {})
    if not analysis:
        flash('No analysis data to generate document. Please upload and analyze files.', 'error')
        return redirect(url_for('upload_files'))
    
    try:
        buffer = generate_word_document(analysis)
        return send_file(
            buffer,
            as_attachment=True,
            download_name='final_document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Document generation failed: {e}")
        flash(f'An error occurred while generating the document: {str(e)}. Please try again.', 'error')
        return redirect(url_for('analysis'))

@app.route('/uploads/<path:filename>')
def serve_uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/remove_files', methods=['POST'])
def remove_files():
    try:
        if os.path.exists(app.config['UPLOAD_FOLDER']):
            # First, remove all files in the directory
            for filename in os.listdir(app.config['UPLOAD_FOLDER']):
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                try:
                    if os.path.isfile(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    print(f"Error removing {file_path}: {e}")
            
            # Then remove the directory itself
            shutil.rmtree(app.config['UPLOAD_FOLDER'])
            
            # Create fresh directory
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            
            # Clear session data
            session.pop('file_paths', None)
            session.pop('analysis', None)
            flash('All uploaded files removed', 'success')
        else:
            flash('No files to remove', 'error')
    except Exception as e:
        print(f"Error removing files: {e}")
        flash('Error removing files', 'error')
    return redirect(url_for('upload_files'))

if __name__ == '__main__':
    app.run(debug=True, use_reloader=False)