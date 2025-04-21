from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_template():
    # Create a new Document
    doc = Document()

    # Set default font and margins
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(12)
    section = doc.sections[0]
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(1)

    # Cover Page
    title = doc.add_paragraph()
    title_run = title.add_run('{PROJECT_TITLE}')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Author: {AUTHOR_NAME}', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Registration Number: {REG_NUMBER}', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Institution & Department: {INSTITUTION}', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Guide: {GUIDE_NAME}', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: {DATE}', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('{TABLE_OF_CONTENTS}', style='Normal')
    doc.add_page_break()

    # Abstract
    heading = doc.add_heading('Abstract', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black headings
    doc.add_paragraph('{ABSTRACT}', style='Normal')

    # Main Report
    doc.add_heading('Main Report', level=1).runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Introduction
    doc.add_heading('Introduction', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph('{INTRODUCTION}', style='Normal')

    # Objectives
    doc.add_heading('Objectives', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph('{OBJECTIVES}', style='Normal')

    # Methodology
    doc.add_heading('Methodology', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph('{METHODOLOGY}', style='Normal')

    # Results & Analysis
    doc.add_heading('Results & Analysis', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph('{ANALYSIS_RESULTS}', style='Normal')

    # Conclusion
    doc.add_heading('Conclusion', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph('{CONCLUSION}', style='Normal')

    # # Future Scope
    # doc.add_heading('Future Scope', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)
    # doc.add_paragraph('{FUTURE_SCOPE}', style='Normal')

    # # Acknowledgement
    # doc.add_heading('Acknowledgement', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)
    # doc.add_paragraph('{ACKNOWLEDGEMENT}', style='Normal')

    # Save the document
    doc.save('sample_template.docx')

if __name__ == '__main__':
    create_sample_template()